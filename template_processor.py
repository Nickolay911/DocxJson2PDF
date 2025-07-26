#!/usr/bin/env python3
"""
Обработчик шаблонов docx с конвертацией в PDF
Автор: Nickolay911
Версия: 1.0
"""

import os
import subprocess
import json
import sys
from docx import Document

def replace_fields_in_docx(docx_path, replacements):
    """
    Заменяет поля в docx документе на значения из словаря
    
    Args:
        docx_path (str): Путь к docx файлу
        replacements (dict): Словарь замен {поле: значение}
    
    Returns:
        Document: Обработанный документ
    """
    doc = Document(docx_path)
    
    # Обработка параграфов
    for paragraph in doc.paragraphs:
        for field, value in replacements.items():
            if field in paragraph.text:
                paragraph.text = paragraph.text.replace(field, value)
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field, value in replacements.items():
                    if field in cell.text:
                        cell.text = cell.text.replace(field, value)
    
    # Обработка заголовков и футеров
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                for field, value in replacements.items():
                    if field in paragraph.text:
                        paragraph.text = paragraph.text.replace(field, value)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for field, value in replacements.items():
                    if field in paragraph.text:
                        paragraph.text = paragraph.text.replace(field, value)
    
    return doc

def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Конвертирует docx файл в PDF используя LibreOffice
    
    Args:
        docx_path (str): Путь к исходному docx файлу
        pdf_path (str): Желаемый путь к PDF файлу
    """
    try:
        # Получаем директорию для PDF
        pdf_dir = os.path.dirname(os.path.abspath(pdf_path))
        if not pdf_dir:
            pdf_dir = "."
        
        # Запускаем LibreOffice для конвертации
        cmd = [
            "libreoffice",
            "--headless", 
            "--convert-to", "pdf",
            "--outdir", pdf_dir,
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            raise Exception(f"Ошибка LibreOffice: {result.stderr}")
        
        # Определяем фактическое имя созданного PDF
        docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
        temp_pdf = os.path.join(pdf_dir, f"{docx_basename}.pdf")
        
        # Переименовываем если нужно
        if temp_pdf != pdf_path and os.path.exists(temp_pdf):
            os.rename(temp_pdf, pdf_path)
        
        if not os.path.exists(pdf_path):
            raise Exception(f"PDF файл не создан: {pdf_path}")
            
    except subprocess.TimeoutExpired:
        raise Exception("Таймаут при конвертации через LibreOffice")
    except FileNotFoundError:
        raise Exception("LibreOffice не найден. Установите: sudo apt install libreoffice")

def load_parameters_from_json(json_path):
    """
    Загружает параметры замены из JSON файла
    
    Args:
        json_path (str): Путь к JSON файлу
        
    Returns:
        dict: Словарь с параметрами замены
    """
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        raise Exception(f"JSON файл не найден: {json_path}")
    except json.JSONDecodeError as e:
        raise Exception(f"Ошибка парсинга JSON: {e}")

def main():
    """
    Основная функция программы
    """
    # Проверка аргументов
    if len(sys.argv) != 4:
        print("Использование: python3 template_processor.py <template.docx> <params.json> <output.pdf>")
        print()
        print("Параметры:")
        print("  template.docx  - входной шаблон docx")
        print("  params.json    - JSON файл с параметрами замены")
        print("  output.pdf     - выходной PDF файл")
        print()
        print("Пример:")
        print("  python3 template_processor.py tpl.docx params.json result.pdf")
        sys.exit(1)
    
    template_path = sys.argv[1]
    json_path = sys.argv[2] 
    pdf_path = sys.argv[3]
    
    # Проверка существования файлов
    if not os.path.exists(template_path):
        print(f"❌ Ошибка: Шаблон не найден: {template_path}")
        sys.exit(1)
        
    if not os.path.exists(json_path):
        print(f"❌ Ошибка: JSON файл не найден: {json_path}")
        sys.exit(1)
    
    try:
        print(f"📄 Обрабатываем шаблон: {template_path}")
        
        # Загружаем параметры
        parameters = load_parameters_from_json(json_path)
        print(f"📋 Загружены параметры: {list(parameters.keys())}")
        
        # Обрабатываем документ
        doc = replace_fields_in_docx(template_path, parameters)
        
        # Создаем временный docx файл
        temp_docx = f"temp_{os.path.basename(template_path)}"
        doc.save(temp_docx)
        print(f"💾 Сохранен обработанный файл: {temp_docx}")
        
        # Конвертируем в PDF
        print(f"🔄 Конвертируем в PDF...")
        convert_docx_to_pdf(temp_docx, pdf_path)
        
        # Удаляем временный файл
        os.remove(temp_docx)
        
        print(f"✅ Готово! PDF создан: {pdf_path}")
        
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main() 